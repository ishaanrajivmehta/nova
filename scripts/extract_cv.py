#!/usr/bin/env python3
"""
extract_cv.py — turn ANY CV into a structured evidence scaffold.

Why this exists: a CV is a compressed artifact. Crafting well needs the raw material a CV throws
away — the shapeable angles behind each bullet, the quantities, the tools actually used. When no
private profile is available, this rebuilds as much of that vault as can be read off the page, and
marks the rest as gaps for the interview step to fill.

It also fixes a real scoring defect. cv_score's quality checks segment bullets heuristically out of
PDF text, and the impact sub-score under-reads verbs through PDF bullet glyphs. Passing the real
bullet list to cv_score.score(bullets=...) bypasses that entirely — so extraction makes the score
more honest, not just the craft.

This is a DETERMINISTIC pre-pass: text, sections, dated role blocks, bullets, tools, contact,
frozen-fact manifest. It deliberately does NOT infer duties, seniority or skills the CV does not
state — inference is the model's job at craft time, under the mode's latitude rules, with the
human confirming. Silent inference here would poison the vault permanently.

CLI:
  python3 extract_cv.py --cv path/to/cv.pdf [--out evidence.json] [--print]
"""
from __future__ import annotations
import argparse, json, os, re, subprocess, sys, tempfile

_HERE = os.path.dirname(os.path.abspath(__file__))
sys.path.insert(0, _HERE)

try:
    import cv_score as _cs
except Exception:
    _cs = None

# --------------------------------------------------------------------------------------------
# Text extraction — pdf / docx / md / txt
# --------------------------------------------------------------------------------------------
def _pdf_text(path):
    try:
        import pdfplumber
        pages = []
        with pdfplumber.open(path) as pdf:
            for p in pdf.pages:
                pages.append(p.extract_text() or "")
        return "\n".join(pages), len(pages)
    except Exception:
        try:
            out = subprocess.run(["pdftotext", "-layout", path, "-"],
                                 capture_output=True, text=True, timeout=60)
            return out.stdout, out.stdout.count("\f") + 1
        except Exception as e:
            raise RuntimeError(f"cannot read PDF ({e}); install pdfplumber or poppler-utils")

def _docx_text(path):
    try:
        import docx  # python-docx
        d = docx.Document(path)
        parts = [p.text for p in d.paragraphs]
        for t in d.tables:
            for row in t.rows:
                parts.extend(c.text for c in row.cells)
        return "\n".join(parts), 1
    except ImportError:
        with tempfile.TemporaryDirectory() as td:
            subprocess.run(["soffice", "--headless", "--convert-to", "txt:Text",
                            "--outdir", td, path], capture_output=True, timeout=120)
            cand = os.path.join(td, os.path.splitext(os.path.basename(path))[0] + ".txt")
            if os.path.exists(cand):
                return open(cand, encoding="utf-8", errors="replace").read(), 1
        raise RuntimeError("cannot read .docx; pip install python-docx or install LibreOffice")

def read_cv(path):
    ext = os.path.splitext(path)[1].lower()
    if ext == ".pdf":
        return _pdf_text(path)
    if ext in (".docx", ".doc", ".odt"):
        return _docx_text(path)
    return open(path, encoding="utf-8", errors="replace").read(), 1

# --------------------------------------------------------------------------------------------
# Structure
# --------------------------------------------------------------------------------------------
SECTION_ALIASES = {
    "summary":    ["summary", "profile", "professional summary", "personal statement", "about me",
                   "career objective", "objective", "personal profile"],
    "experience": ["experience", "work experience", "employment", "employment history",
                   "professional experience", "career history", "work history", "relevant experience"],
    "education":  ["education", "academic background", "qualifications", "academic qualifications",
                   "education and training"],
    "skills":     ["skills", "technical skills", "key skills", "core competencies", "competencies",
                   "expertise", "technical expertise", "skills and expertise"],
    "projects":   ["projects", "personal projects", "selected projects", "portfolio", "key projects"],
    "certs":      ["certifications", "certificates", "licences", "licenses", "professional development",
                   "training", "accreditations"],
    "publications": ["publications", "papers", "research"],
    "awards":     ["awards", "honours", "honors", "achievements"],
    "volunteer":  ["volunteering", "voluntary work", "community"],
}
_SEC_LOOKUP = {a: k for k, v in SECTION_ALIASES.items() for a in v}

# Word exports bullets as PRIVATE USE AREA glyphs ( = Symbol-font bullet), not "•". Missing
# these is why heuristic bullet segmentation under-reads accomplishment lines — and why cv_score's
# impact sub-score reads low on rendered PDFs. Match the whole PUA block, not just typographic marks.
_BULLET = re.compile(r"^\s*(?:[•‣▪●·⁃∙◦▸►*>]|[-]|[-–—](?=\s))\s*")
_MONTH = r"(?:jan|feb|mar|apr|may|jun|jul|aug|sep|sept|oct|nov|dec)[a-z]*\.?"
_DATE_RANGE = re.compile(
    # Feb 2020 – Mar 2021 · 02/2020 – 03/2021 · 2020 – present · and "Apr – Jul 2021" (shared year)
    rf"((?:{_MONTH}\s*)?\d{{4}}|{_MONTH}|\d{{1,2}}/\d{{4}})"
    r"\s*(?:[-–—]|\bto\b)\s*"
    rf"((?:{_MONTH}\s*)?\d{{4}}|\d{{1,2}}/\d{{4}}|present|current|now|date|ongoing)",
    re.I)
_EMAIL = re.compile(r"[\w.+-]+@[\w-]+\.[\w.]+")
_PHONE = re.compile(r"(?:(?:\+\d{1,3}[\s-]?)?(?:\(?\d{2,5}\)?[\s.-]?){2,4}\d{2,4})")
_URL = re.compile(r"(?:https?://|www\.)[^\s,;]+|(?:linkedin\.com|github\.com)/[^\s,;]+", re.I)
_QUANT = re.compile(r"(\d+(?:\.\d+)?\s*%|£\s?\d[\d,.]*\s*[kmb]?|\$\s?\d[\d,.]*\s*[kmb]?|"
                    r"€\s?\d[\d,.]*\s*[kmb]?|\b\d[\d,]{2,}\b|\b\d+\s*(?:x|times)\b|\b\d+\+)", re.I)

def _is_heading(line):
    s = line.strip().strip(":").strip()
    if not s or len(s) > 46 or _BULLET.match(line):
        return None
    key = _SEC_LOOKUP.get(s.lower())
    if key:
        return key
    # ALL-CAPS or Title Case short line that matches an alias loosely
    if s.isupper() or (s.istitle() and len(s.split()) <= 4):
        for alias, k in _SEC_LOOKUP.items():
            if alias in s.lower():
                return k
    return None

def split_sections(text):
    lines = text.splitlines()
    sections, current, buf = {}, "_header", []
    for ln in lines:
        h = _is_heading(ln)
        if h:
            sections.setdefault(current, []).extend(buf)
            current, buf = h, []
        else:
            buf.append(ln)
    sections.setdefault(current, []).extend(buf)
    return {k: [l for l in v] for k, v in sections.items()}

def parse_roles(exp_lines):
    """Group the experience section into dated blocks. Dates are the anchor: they are the one thing
    on a CV that is unambiguous, and they are also a FROZEN FACT that no mode may alter."""
    roles, cur = [], None
    for ln in exp_lines:
        raw = ln.rstrip()
        if not raw.strip():
            continue
        m = _DATE_RANGE.search(raw)
        is_bullet = bool(_BULLET.match(raw))
        if m and not is_bullet:
            if cur:
                roles.append(cur)
            header = _DATE_RANGE.sub("", raw).strip(" |,–-•\t")
            cur = {"header": header, "dates": m.group(0).strip(), "title": None,
                   "employer": None, "bullets": [], "lines": [raw]}
        elif cur is not None:
            cur["lines"].append(raw)
            if is_bullet:
                cur["bullets"].append(_BULLET.sub("", raw).strip())
            elif cur["bullets"]:
                # continuation of a wrapped bullet
                cur["bullets"][-1] = (cur["bullets"][-1] + " " + raw.strip()).strip()
            elif not cur["employer"]:
                cur["employer"] = raw.strip(" |,–-\t")
    if cur:
        roles.append(cur)
    for r in roles:
        parts = re.split(r"\s*[|–—,]\s*|\s+at\s+|\s+•\s*", r["header"], maxsplit=1)
        r["title"] = (parts[0] or "").strip() or None
        if len(parts) > 1 and not r["employer"]:
            r["employer"] = parts[1].strip()
        r["bullets"] = [b for b in r["bullets"] if len(b) > 12]
    return roles

def all_bullets(sections, roles):
    """Every accomplishment line on the CV. This is what gets passed to cv_score(bullets=...)."""
    out = [b for r in roles for b in r["bullets"]]
    for key in ("projects", "volunteer", "_header", "summary"):
        for ln in sections.get(key, []):
            if _BULLET.match(ln):
                t = _BULLET.sub("", ln).strip()
                if len(t) > 12:
                    out.append(t)
    seen, uniq = set(), []
    for b in out:
        k = b.lower()[:60]
        if k not in seen:
            seen.add(k); uniq.append(b)
    return uniq

def detect_skills(text):
    """Canonical skills present in the CV, via the gazetteer + synonym surface forms."""
    if _cs is None:
        return []
    try:
        return sorted(_cs._cv_skill_set(text))
    except Exception:
        return []

def parse_skills_section(sections):
    """Preserve the author's own skill grouping — it is a style signal worth keeping."""
    cats = {}
    for ln in sections.get("skills", []):
        s = _BULLET.sub("", ln).strip()
        if not s:
            continue
        m = re.match(r"^([A-Za-z][A-Za-z /&+-]{2,38}?)\s*[:–-]\s+(.+)$", s)
        if m:
            cats[m.group(1).strip()] = [x.strip() for x in re.split(r"[,;·|]", m.group(2)) if x.strip()]
        else:
            cats.setdefault("_ungrouped", []).extend(
                [x.strip() for x in re.split(r"[,;·|]", s) if x.strip()])
    return cats

def contact_block(text):
    head = "\n".join(text.splitlines()[:12])
    urls = _URL.findall(head)
    phones = [p.strip() for p in _PHONE.findall(head) if len(re.sub(r"\D", "", p)) >= 9]
    return {"emails": _EMAIL.findall(head),
            "phones": phones[:2],
            "links": list(dict.fromkeys(urls))[:5]}

def quality_flags(roles, bullets):
    """What the craft step will need to ask the human about."""
    unquantified = [b for b in bullets if not _QUANT.search(b)]
    weak = []
    if _cs is not None:
        try:
            wl = _cs._data()[1]
            weak = [b for b in bullets
                    if b.split() and b.split()[0].lower() in set(wl.get("weak_verbs", []))]
        except Exception:
            pass
    return {
        "bullet_count": len(bullets),
        "unquantified_bullets": len(unquantified),
        "unquantified_examples": unquantified[:8],
        "weak_opener_bullets": weak[:8],
        "roles_missing_employer": [r["title"] for r in roles if not r["employer"]],
        "roles_with_no_bullets": [r["title"] for r in roles if not r["bullets"]],
    }

def extract(path):
    text, pages = read_cv(path)
    sections = split_sections(text)
    roles = parse_roles(sections.get("experience", []))
    bullets = all_bullets(sections, roles)
    return {
        "source_file": os.path.basename(path),
        "pages": pages,
        "char_count": len(text),
        "contact": contact_block(text),
        "sections_found": [k for k in sections if k != "_header"],
        # FROZEN — no mode, not even loose, may alter anything in here.
        "frozen_facts": {
            "roles": [{"title": r["title"], "employer": r["employer"], "dates": r["dates"]} for r in roles],
            "education_lines": [l.strip() for l in sections.get("education", []) if l.strip()],
            "certification_lines": [l.strip() for l in sections.get("certs", []) if l.strip()],
        },
        "experience": roles,
        "bullets": bullets,
        "skills_declared": parse_skills_section(sections),
        "skills_detected": detect_skills(text),
        "summary_text": " ".join(l.strip() for l in sections.get("summary", []) if l.strip())[:1200],
        "projects_lines": [l.strip() for l in sections.get("projects", []) if l.strip()],
        "flags": quality_flags(roles, bullets),
        "raw_text": text,
    }

def main():
    ap = argparse.ArgumentParser(description="Extract a structured evidence scaffold from a CV")
    ap.add_argument("--cv", required=True)
    ap.add_argument("--out")
    ap.add_argument("--print", action="store_true", dest="show")
    a = ap.parse_args()
    ev = extract(a.cv)
    if a.out:
        slim = {k: v for k, v in ev.items() if k != "raw_text"}
        json.dump(slim, open(a.out, "w", encoding="utf-8"), indent=2, ensure_ascii=False)
        print(f"[written] {a.out}")
    if a.show or not a.out:
        f = ev["frozen_facts"]
        print(f"pages={ev['pages']}  sections={', '.join(ev['sections_found'])}")
        print(f"roles={len(f['roles'])}  bullets={len(ev['bullets'])}  "
              f"skills detected={len(ev['skills_detected'])}")
        for r in f["roles"]:
            print(f"   FROZEN  {str(r['title'])[:38]:38s} | {str(r['employer'])[:26]:26s} | {r['dates']}")
        fl = ev["flags"]
        print(f"   unquantified bullets: {fl['unquantified_bullets']}/{fl['bullet_count']}")
        if fl["weak_opener_bullets"]:
            print(f"   weak openers: {len(fl['weak_opener_bullets'])}")

if __name__ == "__main__":
    main()
